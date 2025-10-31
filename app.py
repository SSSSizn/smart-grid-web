import os
import tempfile
import uuid
from datetime import datetime
import docx
import numpy as np
from flask import Flask, render_template, jsonify, request, send_file, session, redirect, url_for
import pandas as pd
from io import BytesIO
from docx import Document
import matplotlib
import matplotlib.pyplot as plt
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
from io import BytesIO
import pandas as pd
from docx.shared import Pt, Inches
from flask import abort
from werkzeug.utils import secure_filename
from xhtml2pdf import pisa  # 导入xhtml2pdf

matplotlib.use('Agg')  # 设置为非交互式后端

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev_key_for_development_only')  # 用于会话管理

matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 黑体
matplotlib.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 文字保存路径
SAVE_DIR = "saved_texts"
UPLOAD_DIR = "uploads"

os.makedirs(SAVE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)


# HTML 转 PDF 函数
def html_to_pdf(html_content):
    buffer = BytesIO()
    pisa.CreatePDF(html_content, dest=buffer)
    buffer.seek(0)
    return buffer



# 配置上传目录
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "没有文件部分"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "未选择文件"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # 更新全局Excel文件路径
        global EXCEL_FILE
        EXCEL_FILE = filepath

        return jsonify({
            "success": True,
            "message": "文件上传成功",
            "filename": filename
        })

    return jsonify({"error": "不支持的文件类型"}), 400


def get_save_path(page_id):
    return os.path.join(SAVE_DIR, f"{page_id}.txt")


def read_excel(file_path = None, sheet_name='Sheet1'):
    """
    直接读取Excel文件，不需要处理合并单元格
    """
    if file_path is None:
        file_path = EXCEL_FILE

    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError("Excel文件未上传或不存在")

    return pd.read_excel(file_path, sheet_name=sheet_name)


def process_excel_data():
    # 有问题！！ 在导出文件时用到了此函数，但是是针对某个表格写的，要做泛化
    """读取并处理Excel数据，拆分合并单元格，只返回重载或轻载线路"""
    try:
        df = read_excel()

        # 问题线路条件：重载≥90%，轻载≤25%
        problem_conditions = (
                (df['2024年线路最大负载率（%）'] >= 90) |
                (df['2024年线路最大负载率（%）'] <= 25)
        )

        problem_df = df[problem_conditions].copy()
        problem_df = problem_df.sort_values('2024年线路最大负载率（%）', ascending=False)
        problem_df = problem_df.reset_index(drop=True)
        problem_df['序号'] = problem_df.index + 1

        # 空值填充只针对非“重载/轻载原因”列
        problem_df = problem_df.fillna('')

        return problem_df


    except Exception as e:
        print(f"读取文件错误: {e}")
        return pd.DataFrame()



@app.route('/api/problem-lines')
def api_problem_lines():
    # 检查登录状态
    problem_df = process_excel_data()
    result = problem_df.to_dict('records')
    return jsonify({
        'success': True,
        'data': result,
        'total': len(result)
    })


@app.route('/<page_id>')
def dynamic_page(page_id):
    # 检查登录状态
    """根据URL自动匹配对应模板"""
    template_path = os.path.join(app.template_folder, f"{page_id}.html")
    if os.path.exists(template_path):
        return render_template(f"{page_id}.html")
    else:
        return render_template("not_found.html", page_id=page_id), 404


@app.route('/export/<page_id>')
def export_word(page_id):
    """导出单个表格的文件，包含表格、*图表和页面文本内容"""
    # 1. 处理Excel数据生成表格和图表
    df = process_excel_data()

    # 2. 读取页面文本内容
    save_path = get_save_path(page_id)
    try:
        with open(save_path, "r", encoding="utf-8") as f:
            page_text = f.read()
    except FileNotFoundError:
        page_text = "暂无内容。"

    doc = Document()

    # 第一部分：页面文本内容
    doc.add_heading(f"页面 {page_id} 内容", level=1)
    doc.add_paragraph(page_text)

    # 添加分页符（可选）
    # doc.add_page_break()

    # 第二部分：问题线路分析报告（仅当有数据时）
    if not df.empty:
        doc.add_heading("问题线路分析报告", level=1)

        # 插入表格
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            # 设置表头字体大小
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

        # 设置单元格内边距更紧凑
        table.allow_autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1.0)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = docx.oxml.shared.OxmlElement('w:tcMar')
                for side in ['top', 'bottom', 'left', 'right']:
                    mar = docx.oxml.shared.OxmlElement(f'w:{side}')
                    mar.set(docx.oxml.shared.qn('w:w'), '50')
                    tcMar.append(mar)
                tcPr.append(tcMar)

        # 填充数据
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # 保存到内存流
    out_stream = BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)

    return send_file(out_stream, as_attachment=True, download_name=f"{page_id}_表格报告.docx")


@app.route('/export-pdf/<page_id>')
def export_pdf(page_id):
    """导出PDF文件，包含表格、图表和页面文本内容"""
    # 检查登录状态
    if 'user' not in session:
        return redirect(url_for('login'))

    # 获取数据（复用现有逻辑）
    save_path = get_save_path(page_id)
    try:
        with open(save_path, "r", encoding="utf-8") as f:
            page_text = f.read()
    except FileNotFoundError:
        page_text = "暂无内容。"

    df = process_excel_data()  # 数据表格

    # 构建 HTML 内容
    html = f"""
    <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: SimHei, Arial, sans-serif; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .chart {{ margin: 20px 0; text-align: center; }}
            </style>
        </head>
        <body>
            <h1>页面 {page_id} 内容</h1>
            <p>{page_text}</p>
            <h2>问题线路分析报告</h2>
            {df.to_html(index=False)}
        </body>
    </html>
    """

    # 转换为 PDF
    pdf_buffer = html_to_pdf(html)
    return send_file(pdf_buffer, as_attachment=True, download_name=f"{page_id}_报告.pdf", mimetype='application/pdf')


@app.route('/page/<page_id>')
def page(page_id):
    """加载指定页面的HTML"""
    # 读取页面文本内容
    save_path = get_save_path(page_id)
    try:
        with open(save_path, "r", encoding="utf-8") as f:
            text = f.read()
    except FileNotFoundError:
        text = "暂无内容，可以点击编辑并保存。"

    # 直接渲染对应的HTML模板
    return render_template(f"{page_id}.html", text=text, page_id=page_id)


@app.route('/save/<page_id>', methods=['POST'])
def save_text(page_id):
    """保存指定页面的文本"""
    text = request.json.get("text", "")
    with open(get_save_path(page_id), "w", encoding="utf-8") as f:
        f.write(text)
    return jsonify({"status": "ok", "message": f"页面 {page_id} 保存成功！"})


@app.route('/')
def index():
    """首页，列出所有页面链接"""
    if 'user' not in session:
        return redirect(url_for('login'))  # 没登录跳转到登录页

    pages = ['2-5', '3-3']
    return render_template('index.html', pages=pages)



# 登录页面路由
@app.route('/login', methods=['GET', 'POST'])
def login():
    # 如果已登录，跳转到首页
    if 'user' in session:
        return redirect(url_for('index'))

    if request.method == 'POST':
        # 实际应用中应该从数据库验证用户
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')

        # 简单验证，实际应用中应使用更安全的验证方式
        if username == 'admin' and password == 'admin123':  # 仅示例，实际需修改
            session['user'] = username
            return jsonify({'success': True, 'message': '登录成功'})
        else:
            return jsonify({'success': False, 'message': '用户名或密码错误'}), 401

    return render_template('login.html')


# 登出路由
@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))


# 数据上传路由
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}

# @app.route('/upload', methods=['POST'])
# def upload_files():
#     """处理文件上传"""
#     # 检查登录状态
#     if 'user' not in session:
#         return jsonify({'success': False, 'message': '请先登录'}), 401
#
#     # 检查是否有文件被上传
#     if 'files' not in request.files:
#         return jsonify({'success': False, 'message': '没有文件被上传'})
#
#     files = request.files.getlist('files')
#
#     # 处理每个上传的文件
#     for file in files:
#         # 检查文件是否符合要求
#         if file.filename == '':
#             return jsonify({'success': False, 'message': '存在未命名的文件'})
#
#         if file and allowed_file(file.filename):
#             # 确保文件名安全
#             filename = secure_filename(file.filename)
#             # 保存文件到上传目录
#             file_path = os.path.join(UPLOAD_DIR, filename)
#
#             # 如果文件已存在，添加序号避免覆盖
#             counter = 1
#             while os.path.exists(file_path):
#                 name, ext = os.path.splitext(filename)
#                 file_path = os.path.join(UPLOAD_DIR, f"{name}_{counter}{ext}")
#                 counter += 1
#
#             file.save(file_path)
#         else:
#             return jsonify({'success': False, 'message': f'文件 {file.filename} 类型不支持，仅支持 .xlsx 和 .xls 格式'})
#
#     return jsonify({'success': True, 'message': '文件上传成功'})


@app.route('/upload', methods=['GET'])
def upload():
    """显示文件上传页面"""
    # 检查登录状态（与其他路由保持一致的权限控制）
    if 'user' not in session:
        return redirect(url_for('login'))
    # 渲染 upload.html 模板
    return render_template('upload.html')


# 预览页面文字编辑
@app.route('/api/save-table-description', methods=['POST'])
def save_table_description():
    """保存表格描述文字"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    data = request.get_json()
    table_id = data.get('table_id')
    content = data.get('content', '')

    if not table_id:
        return jsonify({'success': False, 'message': '表格ID不能为空'}), 400

    try:
        save_path = get_save_path(table_id)
        with open(save_path, "w", encoding="utf-8") as f:
            f.write(content)

        return jsonify({
            'success': True,
            'message': '描述内容保存成功'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'保存描述内容时出错: {str(e)}'
        }), 500


# 生成报告内容
@app.route('/generate_report', methods=['POST'])
def generate_html_report(selected_tables):
    """生成HTML格式的报告内容"""
    from datetime import datetime
    html_parts = []

    # 报告头部
    html_parts.append('''
    <div class="report-header">
        <h1 class="report-title">智能电网分析报告</h1>
        <div class="report-meta">
            <p>生成时间: {}</p>
            <p>包含表格: {}个</p>
        </div>
    </div>
    '''.format(
        datetime.now().strftime("%Y年%m月%d日 %H:%M:%S"),
        len(selected_tables)
    ))

    # 报告概述
    html_parts.append('''
    <div class="report-section">
        <h2>报告概述</h2>
        <p>本报告基于智能电网系统的各项数据分析结果生成，包含选定的数据表格和分析内容。</p>
    </div>
    ''')

    # 为每个选中的表格添加内容
    for i, table_id in enumerate(selected_tables):
        config = TABLE_CONFIG.get(table_id)
        if not config:
            continue

        # 表格标题
        html_parts.append(f'''
        <div class="table-section" id="table-{table_id}">
            <h3>{config['title']}</h3>
        ''')

        # 添加表格对应的文字内容
        try:
            save_path = get_save_path(table_id)
            with open(save_path, "r", encoding="utf-8") as f:
                text_content = f.read()

            if text_content and text_content != "暂无内容，可以点击编辑并保存。":
                html_parts.append(f'''
                <div class="table-description editable-text" data-table-id="{table_id}">
                    <div class="description-content">{text_content}</div>
                    <div class="edit-controls" style="display: none;">
                        <textarea class="form-control description-edit" rows="4">{text_content}</textarea>
                        <div class="edit-actions mt-2">
                            <button class="btn btn-success btn-sm save-description" data-table-id="{table_id}">保存</button>
                            <button class="btn btn-secondary btn-sm cancel-edit">取消</button>
                        </div>
                    </div>
                    <button class="btn btn-outline-primary btn-sm edit-description-btn">
                        <i class="fas fa-edit"></i> 编辑
                    </button>
                </div>
                ''')
        except FileNotFoundError:
            # 如果没有文字内容，提供一个空的编辑区域
            html_parts.append(f'''
            <div class="table-description editable-text" data-table-id="{table_id}">
                <div class="description-content">暂无内容描述。</div>
                <div class="edit-controls" style="display: none;">
                    <textarea class="form-control description-edit" rows="4" placeholder="请输入表格描述内容..."></textarea>
                    <div class="edit-actions mt-2">
                        <button class="btn btn-success btn-sm save-description" data-table-id="{table_id}">保存</button>
                        <button class="btn btn-secondary btn-sm cancel-edit">取消</button>
                    </div>
                </div>
                <button class="btn btn-outline-primary btn-sm edit-description-btn">
                    <i class="fas fa-edit"></i> 编辑
                </button>
            </div>
            ''')

        # 表格内容
        try:
            df = pd.read_excel(EXCEL_FILE, sheet_name=config['sheet_name'])
            existing_columns = [col for col in config['columns'] if col in df.columns]
            df = df[existing_columns]
            df = df.fillna('')

            # 生成HTML表格
            table_html = df.to_html(
                classes='report-table table table-striped table-bordered',
                index=False,
                escape=False
            )
            html_parts.append(table_html)

            # 添加数据分析摘要
            analysis_summary = generate_analysis_summary(df, table_id)
            html_parts.append(f'''
            <div class="analysis-summary">
                <h4>数据分析摘要</h4>
                {analysis_summary}
            </div>
            ''')

        except Exception as e:
            html_parts.append(f'<div class="alert alert-danger">读取表格数据时出错: {str(e)}</div>')

        html_parts.append('</div>')

        # 添加分页提示
        if i < len(selected_tables) - 1:
            html_parts.append('<div class="page-break"></div>')

    # 报告总结
    html_parts.append('''
    <div class="report-summary">
        <h2>报告总结</h2>
        <p>以上为智能电网分析报告的全部内容。报告基于实际数据生成，反映了当前电网运行的各项指标和状态。</p>
        <div class="signature">
            <p>报告生成系统</p>
            <p>智能电网分析平台</p>
        </div>
    </div>
    ''')

    return ''.join(html_parts)


@app.route('/report_preview', methods=['GET'])
def report_preview():
    """显示报告预览页面"""
    # 检查登录状态（与其他路由保持一致的权限控制）
    if 'user' not in session:
        return redirect(url_for('login'))
    # 渲染 report_preview.html 模板
    return render_template('report_preview.html')


# 导出报告
@app.route('/export_report')
def export_report():
    # 检查登录状态
    if 'user' not in session:
        return redirect(url_for('login'))

    format_type = request.args.get('format', 'word')
    tables = request.args.getlist('tables')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')

    # 生成报告内容
    # 这里应该根据实际表格数据生成内容，与generate_report类似
    report_content = "智能电网分析报告\n\n"
    report_content += "报告概述\n"
    report_content += "本报告包含智能电网系统的各项数据分析结果，基于所选表格数据生成。\n"
    if start_date and end_date:
        report_content += f"数据时间范围: {start_date} 至 {end_date}\n\n"

    # 为了简化，这里只生成文本内容，实际应用中应根据需要生成完整内容

    if format_type == 'pdf':
        # 使用xhtml2pdf生成PDF
        html = f"""
        <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: SimHei, Arial, sans-serif; }}
                    h1, h2 {{ color: #333; }}
                    p {{ margin: 10px 0; }}
                </style>
            </head>
            <body>
                <h1>智能电网分析报告</h1>
                <h2>报告概述</h2>
                <p>本报告包含智能电网系统的各项数据分析结果，基于所选表格数据生成。</p>
                {f'<p>数据时间范围: {start_date} 至 {end_date}</p>' if start_date and end_date else ''}
            </body>
        </html>
        """
        pdf_buffer = html_to_pdf(html)
        return send_file(pdf_buffer, as_attachment=True, download_name='电网分析报告.pdf', mimetype='application/pdf')
    else:
        # 生成Word
        doc = Document()
        doc.add_heading('智能电网分析报告', level=1)

        doc.add_heading('报告概述', level=2)
        doc.add_paragraph('本报告包含智能电网系统的各项数据分析结果，基于所选表格数据生成。')
        if start_date and end_date:
            doc.add_paragraph(f'数据时间范围: {start_date} 至 {end_date}')

        # 保存到内存流
        out_stream = BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)

        return send_file(out_stream, as_attachment=True, download_name='电网分析报告.docx')


# 辅助函数：获取表格标题
def get_table_title(table_id):
    table_titles = {
        'table2-1': '表2-1 xx网格35kV及以上变电站基本情况表',
        'table2-2': '表2-2 xx网格35kV及以上变电站负载情况表',
        'table2-3': '表2-3 低压电网规模一览表',
        'table2-5': '表2-5 中压线路负载率明细表'
    }
    return table_titles.get(table_id, table_id)


# 在文件顶部添加Excel文件路径
EXCEL_FILE = 'data/收资清单填充结果1.xlsx'

# ---定义完整的表格配置字典 ---
# 这是整个系统的核心配置，请务必根据您的实际文档和Excel文件仔细填写
# TABLE_CONFIG = {
#     'table-2-1': {
#         'title': '表2-1 xx网格35kV及以上变电站基本情况表',
#         'sheet_name': '高压变电站基础信息表',  # 对应Excel中的Sheet名
#         'columns': [
#             {'header': '序号', 'source': '序号'},
#             {'header': '变电站名称', 'source': '变电站名称'},
#             {'header': '总容量', 'source': '总容量（MVA）'},
#             {'header': '主变编号', 'source': '主变编号'},
#             {'header': '主变容量', 'source': '主变容量（MVA）'},
#             # 针对“10kV馈线间隔”的精确映射，一个表头对应一个源列
#             {'header': '10kV馈线间隔总数', 'source': '10kV馈线间隔总数'},
#             {'header': '10kV馈线间隔已使用', 'source': '10kV馈线间隔已使用'},
#             {'header': '10kV馈线间隔剩余', 'source': '10kV馈线间隔剩余'},
#             {'header': '未来可扩展间隔数', 'source': '未来可扩展间隔数（个）'},
#             {'header': '备注', 'source': '备注'},
#         ]
#     },
#     'table-2-2': {
#         'title': '表2-2 xx网格35kV及以上变电站负载情况表',
#         'sheet_name': '高压变电站基础信息表',  # 同一个Sheet可以供给多个表格
#         'columns': [
#             {'header': '编号', 'source': '序号'},  # '序号'列在文档中可能被称为'编号'
#             {'header': '变电站名称', 'source': '变电站名称'},
#             {'header': '总容量', 'source': '总容量（MVA）'},
#             {'header': '主变编号', 'source': '主变编号'},
#             {'header': '主变容量', 'source': '主变容量（MVA）'},
#             {'header': '主变典型日负荷', 'source': '主变典型日负荷（MW）'},
#             {'header': '主变典型日负载率', 'source': '主变典型日负载率（%）'},
#             {'header': '主变年最大负荷', 'source': '主变年最大负荷（MW）'},
#             {'header': '主变年最大负载率', 'source': '主变年最大负载率（%）'},
#         ]
#     },
#     'table-2-3': {
#         'title': '表2-3 xx网格供电区域概况表',
#         'sheet_name': '供电区域概况表',
#         'columns': [
#             {'header': '网格', 'source': '网格名称'},
#             {'header': '供区类型', 'source': '供区类型'},
#             {'header': '指标分项', 'source': '指标'},
#             {'header': '数值', 'source': '数值'},
#             {'header': '单位', 'source': '单位'},
#         ]
#     },
#     'table-3-1': {
#         'title': '表3-1 xx网格10kV线路基本情况表',
#         'sheet_name': '10kV线路基础信息表',
#         'columns': [
#             {'header': '线路名称', 'source': '线路名称'},
#             {'header': '起点站', 'source': '起点变电站'},
#             {'header': '联络点', 'source': '主要联络点'},
#             {'header': '线路型号', 'source': '导线型号'},
#             {'header': '线路长度', 'source': '线路长度'},
#             {'header': '开关状态', 'source': '开关状态'},
#         ]
#     },
#     # ... 在此继续添加其他表格的配置
# }
#
# 3333
# # ---  创建新的数据获取路由 ---
# @app.route('/get_table_data/<table_id>')
# def get_table_data(table_id):
#     """
#     根据table_id从配置中获取数据，并返回JSON格式
#     """
#     config = TABLE_CONFIG.get(table_id)
#     if not config:
#         return jsonify({'error': 'Table configuration not found'}), 404
#
#     try:
#         # 1. 读取指定的Excel Sheet
#         df = pd.read_excel(EXCEL_FILE, sheet_name=config['sheet_name'])
#
#         # 2. 提取需要的源列
#         source_columns = [col['source'] for col in config['columns']]
#
#         # 检查所有需要的列是否存在
#         missing_cols = [col for col in source_columns if col not in df.columns]
#         if missing_cols:
#             return jsonify({'error': f'Columns not found in Excel sheet: {", ".join(missing_cols)}'}), 400
#
#         filtered_df = df[source_columns]
#
#         # 3. 重命名列，使其与HTML表头一致
#         # 创建一个从'source'到'header'的映射字典
#         rename_map = {col['source']: col['header'] for col in config['columns']}
#         filtered_df = filtered_df.rename(columns=rename_map)
#
#         # 4. 处理NaN值，避免前端显示问题
#         filtered_df = filtered_df.fillna('')
#
#         # 5. 转换为JSON格式
#         data = filtered_df.to_dict(orient='records')
#         headers = [col['header'] for col in config['columns']]
#
#         return jsonify({
#             'title': config['title'],
#             'headers': headers,
#             'data': data
#         })
#
#     except FileNotFoundError:
#         return jsonify({'error': f'Excel file not found at {EXCEL_FILE}'}), 500
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500
#
#
# 22222
# 添加sheet名称映射字典
SHEET_MAPPING = {
    'table2-1': '1高压变电站基础信息表',
    'table2-2': '2供电区域概况表',
    'table2-3': '3配变负载率明细表',
    'table2-4': '4中压线路负载率表',
    'table2-5': '5高层小区供电保障表',
    'table2-6': '6分布式光伏接入情况表',
    'table2-7': '7充电设施统计表',
    'table2-8': '8台区停电明细表',
    'table2-9': '9台区电压越限统计表',
    'table2-10': '10中压线路网架表',
    'table2-11': '11中压线路分段情况',
    'table2-12': '12中压线路 N-1 校验表',
    'table2-13': '13开关明细表',
    'table2-14': '14配变明细表',
    'table2-15': '15线路规模',
    'table2-16': '16 0.4kV 台区供电半径表',
    'table2-17': '17 10kV线路投运时间表',
    'table2-18': '18线路分线线损情况明细表',
    'table2-19': '19台区线损明细',
    'table2-20': '20 10kV线路配自终端明细表',
    'table2-21': '21线路跨网格供电情况',
    'table2-22': '22线路装接配变容量',
    'table2-23': '23配变情况分布表',
    'table2-24': '24干线型号统计表',
    'table2-25': '25环网柜明细表',
    'table2-26': '26中压线路大分支情况',
    'table2-27': '27中压线路供电半径',
    'table2-28': '28线路状态检测',
    'table2-29': '29问题清单汇总',
    # 添加其他需要的映射...
}


# 添加通用表格路由
# 在app.py顶部添加表格配置字典
TABLE_CONFIG = {
    '2-1': {
        'sheet_name': '1高压变电站基础信息表',
        'columns': ['序号', '变电站名称', '额定容量（MVA）', '主变编号', '10kV馈线间隔总数','10kV馈线间隔已使用', '10kV馈线间隔剩余','未来可扩展间隔数', '备注'],
        'title': '表2-1 xx网格35kV及以上变电站基本情况表'
    },
    '2-2': {
        'sheet_name': '1高压变电站基础信息表',
        'columns': ['变电站名称', '额定容量（MVA）', '主变编号', '主变容量（MVA）', '主变典型日负荷（MW）', '主变典型日负载率（%）', '主变年最大负荷（MW）', '主变年最大负载率（%）'],
        'title': '表2-2 xx网格各变电站主变运行水平表'
    },
    '2-3': {
        'sheet_name': '2供电区域概况表',
        'columns': ['供电网格名称', '供电区域类型', '公变数量（台）', '用户数（万户）','户均配变容量（kVA/户）'],
        'title': '表2-3 低压电网规模一览表'
    },
    '2-4': {
        'sheet_name': '3配变负载率明细表',
        'columns': ['序号', '设备名称', '电压等级', '所属馈线', '额定容量(kVA)', '配变负载率（%）'],
        'title': '表2-4 配变负载率明细表'
    },
    '2-5': {
        'sheet_name': '4中压线路负载率表',
        'columns': ['序号', '变电站名称', '线路名称', '限额电流（A）', '2024年最大电流（A）', '2024年最大负载率（%）', '轻载/重载原因'],  
        'title': '表2-5 中压线路负载率明细表'
    },
    '2-6': {
        'sheet_name': '5高层小区供电保障表',
        'columns': ['序号', '高层小区名称', '供电线路名称', '外线是否双电源', '进线是否双电源'],  
        'title': '表2-6 高层小区明细表'
    },
    '2-7': {
        'sheet_name': '6分布式光伏接入情况表',
        'columns': ['序号', '光伏名称', '台区名称', '电压等级（kV）', '装机容量（kW）', '2024年发电量（kWh）', '2024年上网电量（kWh）', '发电量消纳方式'],  
        'title': '表2-7 xx网格现状分布式光伏接入情况'
    },
    '2-8': {
        'sheet_name': '7充电设施统计表',
        'columns': ['序号', '充电设施名称', '所属线路', '属性1', '属性2', '装机容量（kW）'],  
        'title': '表2-8 充电桩基础设施承载能力'
    },
    '2-9': {
        'sheet_name': '7充电设施统计表',
        'columns': ['序号', '充电设施用户类型', '交流桩总个数', '交流总额定功率（kW）', '直流桩总个数', '直流总额定功率（kW）'],  
        'title': '表2-9 xx网格现状充电基础设施统计表'
    },
    '2-10': {
        'sheet_name': '8台区停电明细表',
        'columns': ['序号', '台区名称', '所属线路', '持续时间（min）', '停电类型', '停电户数', '故障查找时间（h）'],  
        'title': '表2-10 供电可靠性'
    },
    '2-11': {
        'sheet_name': '8台区停电明细表',
        'columns': ['序号', '所属线路', '停电次数', '停电类型（故障停电/计划停电）', '停电原因分析', '持续时间（min）'],  
        'title': '表2-11 频繁停电线路'
    },
    '2-12': {
        'sheet_name': '8台区停电明细表',
        'columns': ['序号', '所属线路', '变电站名称', '停电类型（故障停电/计划停电）', '停电开始时间', '停电终止时间', '持续时间（min）', '停电描述'],  
        'title': '表2-12 xx网格10kV线路停电明细表'
    },
    '2-13': {
        'sheet_name': '台区停电明细表',
        'columns': ['序号', '台区名称', '所属线路', '2024年停电次数', '停电总时长（分钟）', '停电类别（台区自身停电/线路停电引起）', '停电类型（故障停电/计划停电）', '停电原因分析'],  
        'title': '表2-13 xx网格台区停电统计表'
    },
    '2-14': {
        'sheet_name': '9台区电压越限统计表',
        'columns': ['序号', '台区名称', '容量kVA', '台区越上限电压(V)', '越上限时长（h）', '越上限户数', '是否越上限', '是否存在电压三相不平衡'],  
        'title': '表2-14 xx网格台区越上限统计表'
    },
    '2-15': {
        'sheet_name': '9台区电压越限统计表',
        'columns': ['序号', '台区名称', '容量kVA', '台区越上限电压(V)', '越下限时长（h）', '越下限户数', '是否越下限', '是否存在电压三相不平衡'],  
        'title': '表2-15 xx网格台区越下限统计表'
    },
    '2-16': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '线路总条数（条）', '联络率（%）', '站间联络率（%）', '典型接线比例（%）'],  
        'title': '表2-16 xx网格中压线路网架基本情况'
    },
    '2-17': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '变电站名称', '线路名称', '接线方式', '联络线路1', '联络线路2', '是否为标准接线'],  
        'title': '表2-17 中压线路联络情况统计表'
    },
    '2-18': {
        'sheet_name': '11中压线路分段情况',
        'columns': ['序号', '线路名称', '供电半径（km）', '分段数', '分段1容量（kVA）', '分段1台数（台）', '分段2容量（kVA）', '分段2台数（台）', '分段3容量（kVA）', '分段3台数（台）', '分段4容量（kVA）', '分段4台数（台）', '分段5容量（kVA）', '分段5台数（台）', '分段6容量（kVA）', '分段6台数（台）'],  
        'title': '表2-18 中压线路分段情况统计表'
    },
    '2-19': {
        'sheet_name': '11中压线路分段情况',
        'columns': ['序号', '线路名称', '大分支名称', '容量（kVA）', '2024年最大电流值（A）', '备注'],  
        'title': '表2-19 中压线路大分支情况统计表'
    },
    '2-20': {
        'sheet_name': '11中压线路分段情况',
        'columns': ['序号', '变电站', '线路名称', '供电半径（km）'],  
        'title': '表2-20 中压线路供电半径超标明细表'
    },
    '2-21': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '线路名称', '接线方式', '限额电流（A）', '典型日电流（A）', '联络线路1', '线路1转供裕度', '联络线路2', '线路2转供裕度', '最大转供能力（A）', 'N-1校验结果'],  
        'title': '表2-21 中压线路N-1校验结果汇总表'
    },
    '2-22': {
        'sheet_name': '14配变明细表',
        'columns': ['序号', '配变型号', '配变容量', '运行年限', '数量（台）', '占比（%）'],  
        'title': '表2-22 公变情况分布表'
    },
    '2-23': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '供电区域', '总长度（km）', '电缆线路长度（km）', '裸导线长度（km）', '绝缘线长度（km）'],  
        'title': '表2-23 xx网格0.4kV线路规模'
    },
    '2-24': {
        'sheet_name': '16 0.4kV 台区供电半径表',
        'columns': ['序号', '设备名称', '电压等级', '所属馈线', '额定容量（kVA）', '供电半径（m）'],  
        'title': '表2-24 xx网格0.4kV供电半径明细表'
    },
    '2-25': {
        'sheet_name': '17 10kV线路投运时间表',
        'columns': ['序号', '线路名称', '所属变电站', '线路类型', '所属单元', '运行年限（年）'],  
        'title': '表2-25 10kV线路投运时间统计表'
    },
    '2-26': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '线路名称', '所属变电站', '架空线', '电缆线'],  
        'title': '表2-26 中压线路干线型号统计表'
    },
    '2-27': {
        'sheet_name': '10kV线路明细',
        'columns': ['序号', '装接配变容量区间（MVA）', '线路条数（条）', '比例（%）'],  
        'title': '表2-27 线路装接配变容量分布情况'
    },
    '2-28': {
        'sheet_name': '18 线路分线线损情况明细表',
        'columns': ['序号', '变电站名称', '线路名称', '线损率（%）', '备注：是否为高损线路'],  
        'title': '表2-28 中压线路分线线损情况明细表'
    },
    '2-29': {
        'sheet_name': '19 台区线损明细',
        'columns': ['序号', '公用配变总台数', '公用配变线损率平均值（%）', '高损台区台数（台）', '高损台区比例（%）', '负损台区台数（台）', '负损台区比例（%）'],  
        'title': '表2-29 xx网格台区线损情况'
    },
    '2-30': {
        'sheet_name': '20 10kV线路配自终端明细表',
        'columns': ['序号', '名称', '类型', '运行状态', '线路', '在线率', '投退次数', '通信方式', '开关类型', '投运时间'],  
        'title': '表2-30 10kV线路配自终端明细表'
    }
    # 添加其他表格配置...
}

# 修改通用表格路由
@app.route('/table<path:table_id>')
def show_table(table_id):
    config = TABLE_CONFIG.get(table_id)
    if not config:
        return "未找到对应的表格", 404

    try:
        df = pd.read_excel(EXCEL_FILE, sheet_name=config['sheet_name'])
        
        # 读取页面文本内容
        save_path = get_save_path(table_id)
        try:
            with open(save_path, "r", encoding="utf-8") as f:
                text = f.read()
        except FileNotFoundError:
            text = "暂无内容，可以点击编辑并保存。"

        # 只保留配置中指定的列
        existing_columns = [col for col in config['columns'] if col in df.columns]
        df = df[existing_columns]
        df = df.fillna('')
        table_html = df.to_html(classes='table table-striped table-hover', index=False, table_id='data-table')
        return render_template('table_template.html', table_html=table_html, title=config['title'], text=text, page_id=table_id)
    except Exception as e:
        return f"读取表格出错: {str(e)}", 500

# 1111111

@app.route('/api/sheets')
def get_sheets():
    """获取Excel文件中所有sheet的名称"""
    try:
        xls = pd.ExcelFile('data/收资清单填充结果1.xlsx')
        sheet_names = xls.sheet_names
        return jsonify(sheet_names)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/sheet/<sheet_name>')
def get_sheet_data(sheet_name):
    """获取指定sheet的数据"""
    try:
        if not EXCEL_FILE:
            return jsonify({'error': '请先上传Excel文件'}), 400

        df = pd.read_excel(EXCEL_FILE, sheet_name=sheet_name)
        # 处理NaN值
        df = df.fillna('')
        # 转换为JSON格式
        data = df.to_dict('records')
        return jsonify(data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 添加新路由
@app.route('/all_sheets')
def all_sheets():
    """展示所有sheet的页面"""
    return render_template('all_sheets.html')
@app.route('/version1')
def version1():
    """展示所有sheet的页面"""
    return render_template('version1.html')


# === 报告预览相关路由 ===

@app.route('/report')
def report():
    """报告预览页面"""
    if 'user' not in session:
        return redirect(url_for('login'))
    return render_template('report.html')


# 获取所有表格配置的路由
@app.route('/api/table-configs')
def get_table_configs():
    """获取所有表格配置"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    table_configs = []
    for table_id, config in TABLE_CONFIG.items():
        table_configs.append({
            'id': table_id,
            'title': config['title'],
            'sheet_name': config['sheet_name']
        })

    return jsonify({
        'success': True,
        'data': table_configs
    })


@app.route('/api/generate-report-preview', methods=['POST'])
def generate_report_preview():
    """生成报告预览（HTML格式）"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    data = request.get_json()
    selected_tables = data.get('tables', [])

    try:
        # 生成HTML格式的报告预览
        html_content = generate_html_report(selected_tables)

        return jsonify({
            'success': True,
            'html_content': html_content,
            'table_count': len(selected_tables)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'生成报告预览时出错: {str(e)}'
        }), 500


def generate_html_report(selected_tables):
    """生成HTML格式的报告内容"""
    from datetime import datetime
    html_parts = []

    # 报告头部
    html_parts.append('''
    <div class="report-header">
        <h1 class="report-title">智能电网分析报告</h1>
        <div class="report-meta">
            <p>生成时间: {}</p>
            <p>包含表格: {}个</p>
        </div>
    </div>
    '''.format(
        datetime.now().strftime("%Y年%m月%d日 %H:%M:%S"),
        len(selected_tables)
    ))

    # 报告概述
    html_parts.append('''
    <div class="report-section">
        <h2>报告概述</h2>
        <p>本报告基于智能电网系统的各项数据分析结果生成，包含选定的数据表格和分析内容。</p>
    </div>
    ''')

    # 为每个选中的表格添加内容
    for i, table_id in enumerate(selected_tables):
        config = TABLE_CONFIG.get(table_id)
        if not config:
            continue

        # 表格标题
        html_parts.append(f'''
        <div class="table-section" id="table-{table_id}">
            <h3>{config['title']}</h3>
        ''')

        # 表格内容
        try:
            df = pd.read_excel(EXCEL_FILE, sheet_name=config['sheet_name'])
            existing_columns = [col for col in config['columns'] if col in df.columns]
            df = df[existing_columns]
            df = df.fillna('')

            # 生成HTML表格
            table_html = df.to_html(
                classes='report-table table table-striped table-bordered',
                index=False,
                escape=False
            )
            html_parts.append(table_html)

            # 添加数据分析摘要
            analysis_summary = generate_analysis_summary(df, table_id)
            html_parts.append(f'''
            <div class="analysis-summary">
                <h4>数据分析摘要</h4>
                {analysis_summary}
            </div>
            ''')

        except Exception as e:
            html_parts.append(f'<div class="alert alert-danger">读取表格数据时出错: {str(e)}</div>')

        html_parts.append('</div>')

        # 添加分页提示
        if i < len(selected_tables) - 1:
            html_parts.append('<div class="page-break"></div>')

    # 报告总结
    html_parts.append('''
    <div class="report-summary">
        <h2>报告总结</h2>
        <p>以上为智能电网分析报告的全部内容。报告基于实际数据生成，反映了当前电网运行的各项指标和状态。</p>
        <div class="signature">
            <p>报告生成系统</p>
            <p>智能电网分析平台</p>
        </div>
    </div>
    ''')

    return ''.join(html_parts)


def generate_analysis_summary(df, table_id):
    """根据表格数据生成分析摘要"""
    if df.empty:
        return "<p>暂无数据</p>"

    # 基本统计信息
    summary_parts = []
    summary_parts.append(f"<p>数据记录总数: {len(df)} 条</p>")

    # 数值列的统计分析
    numeric_columns = df.select_dtypes(include=[np.number]).columns
    if len(numeric_columns) > 0:
        summary_parts.append("<h5>数值统计</h5>")
        summary_parts.append("<ul>")
        for col in numeric_columns[:3]:  # 只显示前3个数值列
            if col in df.columns:
                mean_val = df[col].mean()
                max_val = df[col].max()
                min_val = df[col].min()
                summary_parts.append(f"<li>{col}: 平均{mean_val:.2f}, 最大{max_val:.2f}, 最小{min_val:.2f}</li>")
        summary_parts.append("</ul>")

    return ''.join(summary_parts)


@app.route('/api/save-report-content', methods=['POST'])
def save_report_content():
    """保存用户编辑的报告内容"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    data = request.get_json()
    html_content = data.get('html_content', '')

    try:
        # 将编辑后的内容保存到session中
        session['edited_report_content'] = html_content

        return jsonify({
            'success': True,
            'message': '报告内容保存成功'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'保存报告内容时出错: {str(e)}'
        }), 500


@app.route('/api/generate-full-report', methods=['POST'])
def generate_full_report():
    """生成包含多个表格的完整报告"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    data = request.get_json()
    selected_tables = data.get('tables', [])

    try:
        # 创建Word文档
        doc = Document()

        # 添加报告标题
        title = doc.add_heading('智能电网分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加报告生成时间
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        time_paragraph = doc.add_paragraph(f"生成时间: {current_time}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 为每个选中的表格添加内容
        for table_id in selected_tables:
            config = TABLE_CONFIG.get(table_id)
            if not config:
                continue

            # 添加表格标题
            doc.add_heading(config['title'], level=1)

            # 添加表格内容
            try:
                df = pd.read_excel(EXCEL_FILE, sheet_name=config['sheet_name'])
                existing_columns = [col for col in config['columns'] if col in df.columns]
                df = df[existing_columns]
                df = df.fillna('')

                # 创建表格
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'

                # 表头
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

                # 数据行
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

            except Exception as e:
                doc.add_paragraph(f"读取表格数据时出错: {str(e)}")

            # 添加分页符（最后一个表格不加）
            if table_id != selected_tables[-1]:
                doc.add_page_break()

        # 保存到内存流
        out_stream = BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)

        # 保存文件到临时目录
        temp_filename = f"report_{uuid.uuid4().hex}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), temp_filename)

        with open(temp_path, 'wb') as f:
            f.write(out_stream.getvalue())

        # 将文件路径存入session，供下载使用
        session['last_report_path'] = temp_path

        return jsonify({
            'success': True,
            'message': '报告生成成功',
            'filename': temp_filename
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'生成报告时出错: {str(e)}'
        }), 500

@app.route('/download-report')
def download_report():
    """下载生成的报告"""
    if 'user' not in session:
        return redirect(url_for('login'))

    file_path = session.get('last_report_path')
    if not file_path or not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '报告文件不存在或已过期'}), 404

    return send_file(file_path, as_attachment=True, download_name='智能电网分析报告.docx')

@app.route('/cleanup-report')
def cleanup_report():
    """清理临时报告文件"""
    if 'user' not in session:
        return jsonify({'success': False, 'message': '未授权'}), 401

    file_path = session.get('last_report_path')
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
            session.pop('last_report_path', None)
        except:
            pass

    return jsonify({'success': True, 'message': '清理完成'})

if __name__ == '__main__':
    app.run(debug=True, port=5001)
